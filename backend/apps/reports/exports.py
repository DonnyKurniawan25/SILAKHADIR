"""Export Laporan Kegiatan ke DOCX & PDF.

Import library berat dilakukan secara lazy di dalam fungsi agar ketidaktersediaan
salah satu dependency tidak membuat seluruh Django gagal startup.
"""
import io
import os

from django.conf import settings


# ----------------------------- Utilities -----------------------------

def _fmt_date(d):
    if not d:
        return '-'
    return d.strftime('%d %B %Y')


def _get_setting():
    try:
        from apps.settings_app.models import AppSetting
        return AppSetting.get_instance()
    except Exception:
        return None


def _absolute_file_url(file_field) -> str:
    """Kembalikan URL publik yang bisa diklik untuk file yang disimpan."""
    if not file_field:
        return ''
    try:
        relative = file_field.url
    except Exception:
        return ''
    base = getattr(settings, 'BACKEND_URL', '').rstrip('/')
    if base:
        return f'{base}{relative}'
    return relative


def _human_size(num_bytes) -> str:
    if num_bytes is None:
        return ''
    try:
        num_bytes = int(num_bytes)
    except Exception:
        return ''
    if num_bytes < 1024:
        return f'{num_bytes} B'
    if num_bytes < 1024 * 1024:
        return f'{num_bytes / 1024:.1f} KB'
    return f'{num_bytes / (1024 * 1024):.1f} MB'


IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg'}


def _is_image_file(filename: str) -> bool:
    if not filename:
        return False
    _, ext = os.path.splitext(filename.lower())
    return ext in IMAGE_EXTENSIONS


# ----------------------------- DOCX -----------------------------

def _html_to_docx_paragraphs(doc, html: str):
    """
    Konversi HTML dari CKEditor ke paragraf DOCX.
    Menangani: p, h1-h4, ul/ol/li, strong/b, em/i, u, br.
    """
    from bs4 import BeautifulSoup  # lazy import

    if not html or not html.strip():
        doc.add_paragraph('—')
        return

    soup = BeautifulSoup(html, 'html.parser')

    def render_runs(paragraph, node, bold=False, italic=False, underline=False):
        if node.name is None:
            text = node.string or ''
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
            return
        new_bold = bold or node.name in ('strong', 'b')
        new_italic = italic or node.name in ('em', 'i')
        new_underline = underline or node.name == 'u'
        if node.name == 'br':
            paragraph.add_run().add_break()
            return
        for child in node.children:
            render_runs(paragraph, child, new_bold, new_italic, new_underline)

    for node in soup.children:
        if getattr(node, 'name', None) is None:
            text = str(node).strip()
            if text:
                doc.add_paragraph(text)
            continue

        if node.name in ('h1', 'h2', 'h3', 'h4'):
            level = int(node.name[1])
            h = doc.add_heading('', level=min(level, 4))
            render_runs(h, node)
        elif node.name == 'p':
            p = doc.add_paragraph()
            render_runs(p, node)
        elif node.name in ('ul', 'ol'):
            style = 'List Bullet' if node.name == 'ul' else 'List Number'
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style=style)
                render_runs(p, li)
        elif node.name == 'br':
            doc.add_paragraph()
        elif node.name in ('blockquote', 'pre'):
            p = doc.add_paragraph(style='Intense Quote')
            render_runs(p, node)
        else:
            p = doc.add_paragraph()
            render_runs(p, node)


def _add_heading_h(doc, text: str, size=14):
    from docx.shared import Pt, RGBColor  # lazy
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x0C, 0x44, 0x69)


def _add_hyperlink(paragraph, url: str, text: str):
    """
    Sisipkan hyperlink ke dalam paragraph python-docx.
    python-docx tidak punya API publik untuk hyperlink, jadi kita
    manipulasi OOXML langsung.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Warna biru + garis bawah
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1685C4')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def build_report_docx(report) -> bytes:
    """Susun DOCX dari EventReport."""
    from docx import Document  # lazy
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Inches, Pt, RGBColor

    event = report.event
    setting = _get_setting()

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # ===== COVER (halaman pertama, full A4) =====
    if report.cover_image and report.cover_image.name:
        try:
            cover_path = report.cover_image.path
            if os.path.exists(cover_path):
                # Gambar cover full halaman (A4 area cetak: ~16cm x 25.7cm)
                cover_p = doc.add_paragraph()
                cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cover_p.paragraph_format.space_before = Pt(0)
                cover_p.paragraph_format.space_after = Pt(0)
                cover_p.add_run().add_picture(
                    cover_path, width=Cm(16), height=Cm(25.7),
                )
                doc.add_page_break()
        except Exception:
            pass

    # Kop
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if setting and setting.institution_logo:
        try:
            logo_path = setting.institution_logo.path
            if os.path.exists(logo_path):
                header.add_run().add_picture(logo_path, width=Inches(0.9))
        except Exception:
            pass

    for text, size, bold in [
        ('PEMERINTAH DAERAH', 11, False),
        (setting.institution_name if setting and setting.institution_name
         else 'Instansi Pemerintah', 14, True),
        (setting.address if setting and setting.address else '', 10, False),
    ]:
        if not text:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after = Pt(12)
    sep_run = sep.add_run('_' * 80)
    sep_run.font.color.rgb = RGBColor(0x0C, 0x44, 0x69)

    # Judul
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('LAPORAN PELAKSANAAN KEGIATAN')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x0C, 0x44, 0x69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(event.title.upper())
    r.bold = True
    r.font.size = Pt(13)

    # A. Identitas
    _add_heading_h(doc, 'A. Identitas Kegiatan')
    info = [
        ('Nama Kegiatan', event.title),
        ('Tema', event.theme or '-'),
        ('Tanggal', f'{_fmt_date(event.start_date)} s.d. {_fmt_date(event.end_date)}'),
        ('Lokasi', event.location or '-'),
        ('Penyelenggara', event.organizer or '-'),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.autofit = True
    for i, (k, v) in enumerate(info):
        row = table.rows[i].cells
        row[0].text = k
        row[1].text = str(v)
        row[0].paragraphs[0].runs[0].bold = True
        row[0].width = Cm(4)

    # B. Ringkasan
    if report.summary:
        _add_heading_h(doc, 'B. Ringkasan Pelaksanaan')
        doc.add_paragraph(report.summary)

    # C. Notulen (HTML)
    _add_heading_h(doc, 'C. Notulen Kegiatan')
    _html_to_docx_paragraphs(doc, report.notulen or '')

    # D. Rekomendasi
    if report.outcome:
        _add_heading_h(doc, 'D. Tindak Lanjut / Rekomendasi')
        doc.add_paragraph(report.outcome)

    # E. Tautan
    links = list(report.links.all())
    if links:
        _add_heading_h(doc, 'E. Tautan Pemberitaan')
        for l in links:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(l.label)
            run.bold = True
            if l.source:
                p.add_run(f' ({l.source})').italic = True
            p.add_run('\n')
            _add_hyperlink(p, l.url, l.url)

    # F. Lampiran berkas - tampilkan isi (gambar embed) + link pendek
    attachments = list(report.attachments.all())
    if attachments:
        _add_heading_h(doc, 'F. Lampiran Berkas')
        doc.add_paragraph('Berkas pendukung kegiatan.')
        for a in attachments:
            url = _absolute_file_url(a.file)
            fname = ''
            try:
                fname = os.path.basename(a.file.name)
            except Exception:
                pass
            size = ''
            try:
                size = _human_size(a.file.size)
            except Exception:
                pass

            # Judul lampiran
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(a.label or fname or 'Lampiran')
            run.bold = True
            if size:
                p.add_run(f'  ({size})').italic = True

            # Jika file gambar, embed langsung di dokumen (ukuran seragam)
            is_img = _is_image_file(fname)
            if is_img:
                try:
                    file_path = a.file.path
                    if os.path.exists(file_path):
                        img_p = doc.add_paragraph()
                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_p.add_run().add_picture(
                            file_path, width=Cm(16), height=Cm(10),
                        )
                except Exception:
                    pass

            # Link pendek (bukan URL panjang)
            if url:
                link_p = doc.add_paragraph()
                link_p.add_run('Unduh: ')
                _add_hyperlink(link_p, url, 'Buka Berkas')

    # G. Foto - 2 foto per halaman, ukuran seragam
    photos = list(report.photos.all())
    if photos:
        _add_heading_h(doc, 'G. Dokumentasi Foto')
        doc.add_paragraph(
            f'Dokumentasi kegiatan ({len(photos)} foto).'
        )

        # Ukuran foto seragam: lebar = lebar area cetak (16cm),
        # tinggi maks = setengah area cetak dikurangi ruang caption (~11cm)
        # Ini memastikan 2 foto muat per halaman.
        PHOTO_WIDTH = Cm(16)
        PHOTO_HEIGHT = Cm(10)

        for i, ph in enumerate(photos):
            try:
                if not ph.image or not os.path.exists(ph.image.path):
                    continue

                # Setiap 2 foto, mulai halaman baru (kecuali foto pertama)
                if i > 0 and i % 2 == 0:
                    doc.add_page_break()

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)

                # Tambah gambar dengan ukuran tetap
                run = p.add_run()
                run.add_picture(ph.image.path, width=PHOTO_WIDTH, height=PHOTO_HEIGHT)

                # Caption di bawah foto
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(12)
                caption_text = ph.caption or f'Foto {i + 1}'
                r = cap.add_run(caption_text)
                r.italic = True
                r.font.size = Pt(9)
            except Exception:
                pass

    # Tanda tangan
    doc.add_paragraph()
    city = (setting.address.split(',')[0].strip() if setting and setting.address else '')
    today = event.end_date.strftime('%d %B %Y') if event.end_date else ''
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.add_run(f'{city}, {today}').italic = True

    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig2.add_run(report.author_position or 'Penyusun').bold = True

    for _ in range(4):
        doc.add_paragraph()

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = name_p.add_run(report.author_name or '_____________________')
    r.bold = True
    r.underline = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------- PDF -----------------------------

PDF_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8" />
<style>
@page {{
  size: A4;
  margin: 2cm 2cm 2cm 2.5cm;
}}
body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; color: #0f172a; line-height: 1.5; }}
.kop {{ text-align: center; border-bottom: 3px solid #0c4469; padding-bottom: 10px; margin-bottom: 18px; }}
.kop img {{ height: 70px; margin-bottom: 4px; }}
.kop .pemda {{ font-size: 10pt; color: #64748b; letter-spacing: 2px; }}
.kop .nama {{ font-size: 15pt; font-weight: bold; color: #0c4469; margin-top: 2px; }}
.kop .alamat {{ font-size: 9pt; color: #64748b; }}
h1.judul {{
  text-align: center; color: #0c4469; font-size: 14pt; margin: 10px 0 4px 0;
  text-transform: uppercase;
}}
h2.subjudul {{ text-align: center; font-size: 12pt; margin: 0 0 18px 0; }}
h3 {{
  color: #0c4469; font-size: 11pt; text-transform: uppercase; letter-spacing: 1px;
  border-bottom: 1px solid #cbd5e1; padding-bottom: 3px; margin-top: 16px; margin-bottom: 8px;
}}
table.info {{ width: 100%; border-collapse: collapse; }}
table.info td {{ padding: 4px 6px; vertical-align: top; font-size: 10pt; }}
table.info td.k {{ width: 130px; font-weight: bold; color: #334155; }}
.notulen p, .notulen li {{ font-size: 10.5pt; }}
.notulen ul, .notulen ol {{ margin-left: 20px; }}
.links ul, .attachments ul {{ padding-left: 18px; }}
.links li, .attachments li {{ margin-bottom: 6px; font-size: 10pt; }}
.links .url, .attachments .url-row a {{ color: #0c4469; font-size: 9pt; word-break: break-all; }}
.attachments .dim {{ color: #64748b; font-size: 9pt; font-weight: normal; }}
.attachments .url-row {{ margin-top: 1px; }}
.photo {{ text-align: center; margin: 8px 0 12px 0; page-break-inside: avoid; }}
.photo img {{ width: 480px; height: 300px; object-fit: cover; border: 1px solid #cbd5e1; }}
.photo .caption {{ font-size: 9pt; font-style: italic; color: #64748b; margin-top: 3px; }}
.ttd {{ margin-top: 30px; text-align: right; font-size: 10pt; }}
.ttd .gap {{ height: 60px; }}
.ttd .nama {{ font-weight: bold; text-decoration: underline; }}
.cover-page {{ text-align: center; page-break-after: always; }}
.cover-page img {{ width: 100%; height: 100%; object-fit: contain; }}
</style>
</head>
<body>

{cover_block}

<div class="kop">
  {logo_src_html}
  <div class="pemda">PEMERINTAH DAERAH</div>
  <div class="nama">{institution_name}</div>
  <div class="alamat">{address}</div>
</div>

<h1 class="judul">Laporan Pelaksanaan Kegiatan</h1>
<h2 class="subjudul">{event_title}</h2>

<h3>A. Identitas Kegiatan</h3>
<table class="info">
  <tr><td class="k">Nama Kegiatan</td><td>{event_title}</td></tr>
  <tr><td class="k">Tema</td><td>{event_theme}</td></tr>
  <tr><td class="k">Tanggal</td><td>{event_date}</td></tr>
  <tr><td class="k">Lokasi</td><td>{event_location}</td></tr>
  <tr><td class="k">Penyelenggara</td><td>{event_organizer}</td></tr>
</table>

<h3>B. Ringkasan Pelaksanaan</h3>
<p>{summary}</p>

<h3>C. Notulen Kegiatan</h3>
<div class="notulen">{notulen_html}</div>

<h3>D. Tindak Lanjut / Rekomendasi</h3>
<p>{outcome}</p>

{links_block}
{attachments_block}
{photos_block}

<div class="ttd">
  <div>{city}{today_sep}{today}</div>
  <div>{author_position}</div>
  <div class="gap"></div>
  <div class="nama">{author_name}</div>
</div>

</body>
</html>
"""


def build_report_pdf(report) -> bytes:
    """Generate PDF laporan via xhtml2pdf dari template HTML."""
    from xhtml2pdf import pisa  # lazy import

    event = report.event
    setting = _get_setting()

    logo_src_html = ''
    if setting and setting.institution_logo:
        try:
            path = setting.institution_logo.path
            if os.path.exists(path):
                logo_src_html = f'<img src="{path.replace(chr(92), "/")}" />'
        except Exception:
            pass

    # Cover halaman pertama (full A4)
    cover_block = ''
    if report.cover_image and report.cover_image.name:
        try:
            cover_path = report.cover_image.path
            if os.path.exists(cover_path):
                src = cover_path.replace('\\', '/')
                cover_block = f'<div class="cover-page"><img src="{src}" /></div>'
        except Exception:
            pass

    photos_html = ''
    photo_list = list(report.photos.all())
    for i, ph in enumerate(photo_list):
        try:
            if ph.image and os.path.exists(ph.image.path):
                src = ph.image.path.replace('\\', '/')
                cap = (ph.caption or f'Foto {i + 1}').replace('<', '&lt;')
                # Page break setiap 2 foto (kecuali awal)
                page_break = '<div style="page-break-before:always;"></div>' if (i > 0 and i % 2 == 0) else ''
                photos_html += (
                    f'{page_break}'
                    f'<div class="photo"><img src="{src}" />'
                    f'<div class="caption">{cap}</div></div>'
                )
        except Exception:
            continue

    links_html = ''
    for l in report.links.all():
        label = (l.label or '').replace('<', '&lt;')
        src = f' <i>({l.source})</i>' if l.source else ''
        url = (l.url or '').replace('<', '&lt;')
        links_html += f'<li><b>{label}</b>{src}<br/><span class="url">{url}</span></li>'

    attachments_html = ''
    for a in report.attachments.all():
        label = (a.label or '').replace('<', '&lt;')
        fname = ''
        try:
            fname = os.path.basename(a.file.name)
        except Exception:
            pass
        size = _human_size(getattr(a.file, 'size', None))
        url = _absolute_file_url(a.file)
        is_img = _is_image_file(fname)

        # Embed gambar langsung jika file gambar
        img_tag = ''
        if is_img:
            try:
                file_path = a.file.path.replace('\\', '/')
                if os.path.exists(a.file.path):
                    img_tag = (
                        f'<div style="text-align:center; margin:6px 0;">'
                        f'<img src="{file_path}" style="width:480px; height:300px; object-fit:cover; border:1px solid #cbd5e1;" />'
                        f'</div>'
                    )
            except Exception:
                pass

        size_str = f' ({size})' if size else ''
        link_tag = f'<div class="url-row"><a href="{url}">Buka Berkas</a></div>' if url else ''
        attachments_html += (
            f'<li><b>{label or fname or "Lampiran"}</b>{size_str}'
            f'{img_tag}{link_tag}</li>'
        )

    city = (setting.address.split(',')[0].strip() if setting and setting.address else '')
    today = event.end_date.strftime('%d %B %Y') if event.end_date else ''

    ctx = {
        'cover_block': cover_block,
        'institution_name': (setting.institution_name if setting and setting.institution_name
                             else 'Instansi Pemerintah'),
        'address': setting.address if setting else '',
        'logo_src_html': logo_src_html,
        'event_title': event.title,
        'event_theme': event.theme or '-',
        'event_date': f'{_fmt_date(event.start_date)} s.d. {_fmt_date(event.end_date)}',
        'event_location': event.location or '-',
        'event_organizer': event.organizer or '-',
        'summary': (report.summary or '—').replace('\n', '<br/>'),
        'notulen_html': report.notulen or '<p>—</p>',
        'outcome': (report.outcome or '—').replace('\n', '<br/>'),
        'author_name': report.author_name or '_____________________',
        'author_position': report.author_position or 'Penyusun',
        'links_block': (f'<h3>E. Tautan Pemberitaan</h3><div class="links"><ul>{links_html}</ul></div>'
                        if links_html else ''),
        'attachments_block': (
            f'<h3>F. Lampiran Berkas</h3>'
            f'<p style="margin-top:-4px; font-size:10pt; color:#334155;">'
            f'Berkas pendukung kegiatan.'
            f'</p>'
            f'<div class="attachments"><ul>{attachments_html}</ul></div>'
            if attachments_html else ''
        ),
        'photos_block': (f'<h3>G. Dokumentasi Foto</h3>{photos_html}'
                         if photos_html else ''),
        'city': city,
        'today_sep': ', ' if (city and today) else '',
        'today': today,
    }

    html = PDF_TEMPLATE.format(**ctx)
    buf = io.BytesIO()
    pisa_status = pisa.CreatePDF(src=html, dest=buf, encoding='utf-8')
    if pisa_status.err:
        raise RuntimeError(f'Gagal membuat PDF: {pisa_status.err}')
    return buf.getvalue()
